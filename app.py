import streamlit as st
import pandas as pd
from datetime import datetime
import os
from PIL import Image
import io
from docx import Document
from docx.shared import Inches

# 页面配置
st.set_page_config(page_title="生活看板Pro", layout="wide")

# 文件夹准备
IMAGE_DIR = "health_images"
if not os.path.exists(IMAGE_DIR):
    os.makedirs(IMAGE_DIR)

DATA_FILE = "health_records.csv"

# --- 核心逻辑：数据加载 ---
def load_data():
    if os.path.exists(DATA_FILE):
        try:
            df = pd.read_csv(DATA_FILE)
            df['日期'] = df['日期'].astype(str)
            return df.drop_duplicates(subset=['日期'], keep='last').reset_index(drop=True)
        except:
            return pd.DataFrame()
    return pd.DataFrame()

if 'data' not in st.session_state:
    st.session_state.data = load_data()

# --- 图像处理 ---
def process_and_save(img_file, slot_name, date_str):
    if img_file is None:
        return None  
    try:
        img = Image.open(img_file)
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        file_name = f"{date_str}_{slot_name}.jpg"
        save_path = os.path.join(IMAGE_DIR, file_name)
        img.save(save_path, "JPEG", quality=40, optimize=True) 
        return save_path
    except Exception as e:
        st.error(f"图片保存失败: {e}")
        return None

# --- Word 导出逻辑 (图文并茂) ---
def export_to_word(selected_df):
    doc = Document()
    doc.add_heading('生活看板 - 历史健康报告', 0)
    
    for _, row in selected_df.iterrows():
        doc.add_heading(f"日期: {row['日期']}", level=1)
        
        # 基础指标表格
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = f"体重: {row['体重']}kg"
        hdr_cells[1].text = f"抽烟: {row['抽烟']}根"
        hdr_cells[2].text = f"用药: {row['用药按时']}"
        
        doc.add_paragraph(f"😴 睡眠: {row['昨日入睡']} ~ {row['今早起床']} (时长: {row['睡眠时长']}h)")
        doc.add_paragraph(f"🚽 生理: 起夜: {row['是否起夜']} | 排便性状: {row['排便性状']}")
        doc.add_paragraph(f"🔋 状态: 精力: {row['精力']} | 情绪: {row['情绪']}")
        doc.add_paragraph(f"🚶 运动: 午后步: {row['午后步行']}min | 晚后步: {row['晚后步行']}min")
        
        # 饮食明细
        doc.add_heading('🍲 饮食记录', level=2)
        p = doc.add_paragraph()
        p.add_run(f"早餐: {row['早餐']}\n").bold = True
        p.add_run(f"午餐: {row['午餐']}\n").bold = True
        p.add_run(f"加餐: {row['加餐']}\n").bold = True
        p.add_run(f"晚餐: {row['晚餐']}").bold = True
        
        # 插入图片
        doc.add_heading('📸 饮食照片', level=2)
        for label, img_key in [("早餐", "早餐图"), ("午餐", "午餐图"), ("加餐", "加餐图"), ("晚餐", "晚餐图")]:
            img_path = str(row[img_key])
            if img_path != "na" and os.path.exists(img_path):
                doc.add_paragraph(f"[{label}图片]")
                doc.add_picture(img_path, width=Inches(3.0))
        
        doc.add_paragraph(f"📝 今日回顾: {row['回顾']}")
        doc.add_page_break()

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- UI 渲染 ---
st.title("🍎 个人健康看板 Pro")

tab1, tab2 = st.tabs(["✍️ 详细登记", "📊 历史管理"])

with tab1:
    curr_date = st.date_input("选择记录日期", value=datetime.now())
    date_str = str(curr_date)
    
    existing_df = st.session_state.data
    record = existing_df[existing_df['日期'] == date_str].iloc[0] if not existing_df.empty and date_str in existing_df['日期'].values else None

    with st.form("main_form"):
        st.subheader("1. 基础指标")
        c1, c2, c3 = st.columns(3)
        with c1:
            weight = st.number_input("体重 (kg)", value=float(record['体重']) if record is not None else 84.25, step=0.01)
            smoke = st.number_input("抽烟 (根)", value=int(record['抽烟']) if record is not None else 10, step=1)
        with c2:
            s_in = st.text_input("昨日入睡", value=record['昨日入睡'] if record is not None else "1:20")
            s_out = st.text_input("今早起床", value=record['今早起床'] if record is not None else "8:40")
            s_dur = st.number_input("睡眠时长 (h)", value=float(record['睡眠时长']) if record is not None else 8.0, step=0.5)
        with c3:
            s_night = st.selectbox("是否起夜", ["否", "是"], index=0 if record is None or record['是否起夜']=="否" else 1)
            bowel = st.selectbox("排便性状", ["1","2","3","4","5","6","7"], index=3 if record is None else int(record['排便性状'])-1)
            meds = st.selectbox("用药/补剂按时", ["是", "否"], index=0 if record is None or record['用药按时']=="是" else 1)

        st.divider()
        st.subheader("2. 饮食与运动")
        c4, c5 = st.columns(2)
        with c4:
            water = st.number_input("饮水 (L)", value=float(record['饮水']) if record is not None else 2.0, step=0.1)
            water_extra = st.text_input("饮料", value=record['其中饮料'] if record is not None else "na")
            breakfast_t = st.text_area("早餐及时间", value=record['早餐'] if record is not None else "")
            lunch_t = st.text_area("午餐及时间", value=record['午餐'] if record is not None else "")
        with c5:
            snack_t = st.text_area("加餐", value=record['加餐'] if record is not None else "na")
            dinner_t = st.text_area("晚餐及时间", value=record['晚餐'] if record is not None else "")
            walk_l = st.number_input("午后步行(min)", value=int(record['午后步行']) if record is not None else 0)
            walk_d = st.number_input("晚后步行(min)", value=int(record['晚后步行']) if record is not None else 0)

        st.divider()
        st.subheader("3. 状态与回顾")
        c6, c7 = st.columns(2)
        with c6:
            energy = st.select_slider("精力", options=["低", "中", "高"], value=record['精力'] if record is not None else "中")
            mood = st.select_slider("情绪", options=["低", "中", "高"], value=record['情绪'] if record is not None else "中")
        with c7:
            discomfort = st.text_area("不适/特殊", value=record['不适'] if record is not None else "无")
            review = st.text_area("回顾", value=record['回顾'] if record is not None else "")

        if st.form_submit_button("💾 暂存当前文字信息"):
            st.toast(f"文字信息已暂存 (日期: {date_str})", icon="ℹ️")
            st.success("暂存完成！别忘了在下方同步照片或点击同步数据库。")

    st.divider()
    st.subheader("4. 照片管理")
    ci1, ci2 = st.columns(2)
    def image_slot(label, key_prefix):
        st.write(f"**{label}**")
        cam = st.camera_input(f"拍{label}", key=f"{key_prefix}_cam_{date_str}")
        up = st.file_uploader(f"传{label}", type=['jpg','png','jpeg'], key=f"{key_prefix}_up_{date_str}")
        img = cam if cam else up
        if img and st.button(f"❌ 清除{label}", key=f"clr_{key_prefix}_{date_str}"):
            st.rerun()
        return img

    with ci1:
        img_b = image_slot("早餐", "b")
        img_l = image_slot("午餐", "l")
    with ci2:
        img_s = image_slot("加餐", "s")
        img_d = image_slot("晚餐", "d")

    if st.button("🚀 实时同步到数据库 (历史管理)", type="primary", use_container_width=True):
        new_row = {
            "日期": date_str, "昨日入睡": s_in, "今早起床": s_out, "睡眠时长": s_dur,
            "是否起夜": s_night, "排便性状": bowel, "体重": weight, "用药按时": meds,
            "抽烟": smoke, "饮水": water, "其中饮料": water_extra,
            "早餐": breakfast_t, "午餐": lunch_t, "午后步行": walk_l,
            "加餐": snack_t, "晚餐": dinner_t, "晚后步行": walk_d,
            "精力": energy, "情绪": mood, "不适": discomfort, "回顾": review
        }
        for slot, img, tag in [("早餐图", img_b, "breakfast"), ("午餐图", img_l, "lunch"), 
                              ("加餐图", img_s, "snack"), ("晚餐图", img_d, "dinner")]:
            path = process_and_save(img, tag, date_str)
            new_row[slot] = path if path else (record[slot] if record is not None else "na")

        df = st.session_state.data
        if not df.empty and date_str in df['日期'].values:
            df.loc[df['日期'] == date_str, new_row.keys()] = new_row.values()
        else:
            df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)

        df.to_csv(DATA_FILE, index=False)
        st.session_state.data = df
        st.balloons()
        st.success(f"🎉 同步已完成！{date_str} 的记录已存入历史管理。")

with tab2:
    st.subheader("📜 历史数据全量管理")
    df_view = st.session_state.data
    if not df_view.empty:
        df_view = df_view.sort_values("日期", ascending=False)
        
        # 导出区域
        st.write("### 📤 导出报告")
        sel_dates = st.multiselect("勾选日期导出 Word：", options=df_view['日期'].tolist())
        if sel_dates:
            word_data = export_to_word(df_view[df_view['日期'].isin(sel_dates)])
            st.download_button("📥 下载 Word 图文报告", data=word_data, file_name=f"健康报告_{datetime.now().strftime('%m%d')}.docx")
        st.divider()

        # 展示区域
        for _, row in df_view.iterrows():
            with st.expander(f"📅 {row['日期']} | 体重: {row['体重']}kg | 睡眠: {row['睡眠时长']}h"):
                col_info1, col_info2 = st.columns(2)
                with col_info1:
                    st.markdown(f"**😴 睡眠**: {row['昨日入睡']} - {row['今早起床']}")
                    st.markdown(f"**🚽 生理**: 起夜:{row['是否起夜']} | 便性:{row['排便性状']} | 用药:{row['用药按时']}")
                    st.markdown(f"**🚬 生活**: 抽烟:{row['抽烟']}根 | 饮水:{row['饮水']}L ({row['其中饮料']})")
                with col_info2:
                    st.markdown(f"**⚡ 状态**: 精力:{row['精力']} | 情绪:{row['情绪']}")
                    st.markdown(f"**🚶 步行**: 午后:{row['午后步行']}min | 晚后:{row['晚后步行']}min")
                    st.markdown(f"**⚠️ 不适**: {row['不适']}")
                
                st.info(f"📝 **今日回顾**: {row['回顾']}")
                
                # 饮食详情
                st.write("**🍱 饮食图文**")
                m1, m2, m3, m4 = st.columns(4)
                meals = [("早餐", "早餐图", m1), ("午餐", "午餐图", m2), ("加餐", "加餐图", m3), ("晚餐", "晚餐图", m4)]
                for label, key, col in meals:
                    with col:
                        st.caption(f"{label}: {row[label]}")
                        if str(row[key]) != "na" and os.path.exists(str(row[key])):
                            st.image(str(row[key]), use_container_width=True)
                
                if st.button("🗑️ 删除此条", key=f"del_{row['日期']}"):
                    st.session_state.data = st.session_state.data[st.session_state.data['日期'] != row['日期']]
                    st.session_state.data.to_csv(DATA_FILE, index=False)
                    st.rerun()
    else:
        st.info("暂无历史记录。")
